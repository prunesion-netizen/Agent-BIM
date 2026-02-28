"""
bim_generators.py — Generatoare documente BIM pentru orice proiect
Foloseste RAG + Claude AI pentru a genera documente profesionale DOCX.

Utilizare din Flask:
    from bim_generators import generate_document
    path = generate_document("bep", "Ghizela")
"""

import os
import re
import datetime
import threading
from pathlib import Path

import json

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from anthropic import Anthropic
from dotenv import load_dotenv

load_dotenv()

# ── Config ─────────────────────────────────────────────────────────────────────
GENERATED_DIR = Path("generated")
GENERATED_DIR.mkdir(exist_ok=True)

AI_MODEL     = "claude-sonnet-4-6"
BLUE_DARK    = RGBColor(0x1D, 0x4E, 0xD8)
GRAY_TEXT    = RGBColor(0x37, 0x41, 0x51)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

_ai_client = None
def get_ai():
    global _ai_client
    if _ai_client is None:
        _ai_client = Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
    return _ai_client

# ── RAG helper ─────────────────────────────────────────────────────────────────
try:
    from bim_rag import query_rag as _query_rag
    _RAG_OK = True
except Exception:
    _RAG_OK = False
    def _query_rag(q, n=5):
        return {"context": "", "sources": [], "rag_used": False}


def get_project_context(project: str, queries: list, n: int = 5) -> str:
    """Extrage context RAG pentru un proiect și o lista de interogari."""
    parts = []
    for q in queries:
        full_q = f"{q} {project}" if project else q
        r = _query_rag(full_q, n=n)
        if r.get("rag_used") and r.get("context"):
            parts.append(r["context"])
    return "\n\n---\n\n".join(parts)


def ask_claude(system: str, user: str, max_tokens: int = 3000) -> str:
    resp = get_ai().messages.create(
        model=AI_MODEL,
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": user}],
    )
    return resp.content[0].text.strip()


# ── DOCX Helpers ───────────────────────────────────────────────────────────────
def _shd(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _setup_doc(title: str, project: str) -> Document:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.0)
    # Cover header
    p = doc.add_paragraph()
    r = p.add_run("━" * 55)
    r.font.color.rgb = BLUE_DARK
    r.font.size = Pt(12)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(title)
    r2.font.size = Pt(20)
    r2.font.bold = True
    r2.font.color.rgb = BLUE_DARK

    if project:
        p3 = doc.add_paragraph()
        r3 = p3.add_run(f"Proiect: {project}")
        r3.font.size = Pt(13)
        r3.font.color.rgb = GRAY_TEXT

    p4 = doc.add_paragraph()
    r4 = p4.add_run(f"Generat: {datetime.date.today().strftime('%d.%m.%Y')} · Agent BIM Romania")
    r4.font.size = Pt(9)
    r4.font.italic = True
    r4.font.color.rgb = GRAY_TEXT
    doc.add_paragraph()
    return doc


def _add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    if p.runs:
        p.runs[0].font.color.rgb = BLUE_DARK
        p.runs[0].font.size = Pt(15)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)


def _add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    if p.runs:
        p.runs[0].font.color.rgb = GRAY_TEXT
        p.runs[0].font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)


def _add_body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY_TEXT
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY_TEXT
    return p


def _add_table(doc, headers: list, rows: list):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hr = t.rows[0]
    for j, h in enumerate(headers):
        _shd(hr.cells[j], "1D4ED8")
        p = hr.cells[j].paragraphs[0]
        r = p.add_run(str(h))
        r.font.bold  = True
        r.font.size  = Pt(9)
        r.font.color.rgb = WHITE
        p.alignment  = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for i, row in enumerate(rows):
        dr = t.rows[i + 1]
        bg = "F0F9FF" if i % 2 == 0 else "FFFFFF"
        for j, cell in enumerate(row):
            _shd(dr.cells[j], bg)
            p = dr.cells[j].paragraphs[0]
            r = p.add_run(str(cell))
            r.font.size = Pt(8.5)
            r.font.color.rgb = GRAY_TEXT
            if j == 0:
                r.font.bold = True
    doc.add_paragraph()


def _set_cell_border(cell, **kwargs):
    """Adaugă chenar la o celulă."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), kwargs.get("val", "single"))
        border.set(qn("w:sz"), str(kwargs.get("sz", 4)))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), kwargs.get("color", "1D4ED8"))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_info_table(doc, rows: list):
    """Tabel cheie-valoare pentru informații proiect (2 coloane: cheie albastru, valoare gri)."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [Cm(5.5), Cm(11.5)]
    for i, (key, val) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        _shd(row.cells[0], "DBF0FE")
        kp = row.cells[0].paragraphs[0]
        kr = kp.add_run(key)
        kr.font.bold = True
        kr.font.size = Pt(9)
        kr.font.color.rgb = BLUE_DARK
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(str(val))
        vr.font.size = Pt(9)
        vr.font.color.rgb = GRAY_TEXT
    doc.add_paragraph()


def _add_matrix_table(doc, headers: list, data: list):
    """Tabel matrice cu header albastru și rânduri alternate."""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = "Table Grid"
    hrow = table.rows[0]
    for j, h in enumerate(headers):
        _shd(hrow.cells[j], "1D4ED8")
        p = hrow.cells[j].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row_data in enumerate(data):
        drow = table.rows[i + 1]
        bg = "F0F9FF" if i % 2 == 0 else "FFFFFF"
        for j, cell_text in enumerate(row_data):
            _shd(drow.cells[j], bg)
            p = drow.cells[j].paragraphs[0]
            r = p.add_run(str(cell_text))
            r.font.size = Pt(8.5)
            r.font.color.rgb = GRAY_TEXT
            if j == 0:
                r.font.bold = True
    doc.add_paragraph()


def _md_to_doc(doc, text: str):
    """Conversie Markdown simplu → paragrafele DOCX (heading, bullet, body)."""
    lines = text.split("\n")
    for line in lines:
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("### "):
            _add_h2(doc, line[4:])
        elif line.startswith("## "):
            _add_h1(doc, line[3:])
        elif line.startswith("# "):
            _add_h1(doc, line[2:])
        elif re.match(r"^[-*•]\s+", line):
            _add_bullet(doc, re.sub(r"^[-*•]\s+", "", line))
        elif re.match(r"^\d+\.\s+", line):
            _add_bullet(doc, re.sub(r"^\d+\.\s+", "", line))
        else:
            _add_body(doc, line)


# ══════════════════════════════════════════════════════════════════════════════
# GENERATOARE
# ══════════════════════════════════════════════════════════════════════════════

SYSTEM_BIM = (
    "Ești expert BIM senior în România, specializat în ISO 19650, BEP, EIR, LOD, CDE. "
    "Generezi documente profesionale, complete, în română. "
    "Folosești datele din context când există; altfel folosești standardele internaționale. "
    "Formatul răspunsului: Markdown cu ## pentru capitole, ### pentru subcapitole, - pentru liste."
)


def gen_bep(project: str) -> str:
    """Generează BIM Execution Plan. Returnează calea fișierului."""
    ctx = get_project_context(project, [
        "obiectul contractului beneficiar proiectant",
        "cerinte BIM modelare software livrabile",
        "Common Data Environment CDE platforma",
        "LOD nivel detaliu faze proiect",
        "standarde ISO 19650 BIM",
        "echipa BIM manager coordinator responsabilitati",
    ])

    prompt = f"""Genereaza un BEP (BIM Execution Plan) complet pentru proiectul: **{project}**

Context din documentele proiectului:
{ctx or "Nu exista context specific – foloseste standarde generale BIM si ISO 19650."}

Include TOATE sectiunile standard:
## 1. Informatii generale proiect
## 2. Obiectivele BIM (OIR / AIR / PIR)
## 3. Echipa BIM – Roluri si Responsabilitati (RACI)
## 4. Mediul Comun de Date (CDE)
## 5. Standarde, Software si Formate
## 6. Nivelele de Detaliu (LOD 100-400) pe faze
## 7. Livrabile BIM si Calendar
## 8. Coordonare si Clash Detection
## 9. Documentatie As-Built si Predare
## 10. Managementul Calitatii BIM
## 11. Aprobari si Istoricul Reviziilor

Fii detaliat, profesional. Minim 800 cuvinte."""

    content = ask_claude(SYSTEM_BIM, prompt, max_tokens=4000)

    doc = _setup_doc("PLAN DE EXECUȚIE BIM (BEP)", project)
    _md_to_doc(doc, content)

    ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    safe = re.sub(r"[^\w]", "_", project)[:30]
    out  = GENERATED_DIR / f"BEP_{safe}_{ts}.docx"
    doc.save(str(out))
    return str(out)


def gen_lod(project: str) -> str:
    """Generează Matrice LOD. Returnează calea fișierului."""
    ctx = get_project_context(project, [
        "faze proiect PT DDE executie receptie",
        "elemente BIM structuri instalatii arhitectura",
        "LOD nivel detaliu livrabile",
    ])

    prompt = f"""Genereaza o Matrice LOD completa pentru proiectul: **{project}**

Context:
{ctx or "Proiect de constructii – foloseste matrice LOD standard."}

## 1. Introducere – Ce este LOD
## 2. Definitii LOD (100 / 200 / 300 / 350 / 400)
## 3. Matrice LOD pe elemente si faze
(descrie textual: Element | PT | DDE | Executie | As-Built | Note)
## 4. Responsabilitati LOD pe rol
## 5. Verificare si validare LOD

Fii specific. Include minimum 15 tipuri de elemente BIM relevante pentru proiect."""

    content = ask_claude(SYSTEM_BIM, prompt, max_tokens=3000)

    doc = _setup_doc("MATRICE LOD / LOI", project)
    _md_to_doc(doc, content)

    # Adauga tabelul LOD standard
    _add_h1(doc, "Tabel de referință LOD")
    _add_table(doc,
        ["LOD", "Denumire", "Precizie geometrie", "Informații incluse"],
        [
            ["100", "Concept",     "Simbolic / volum",          "Suprafețe, volume, orientare aproximativă"],
            ["200", "Schematic",   "Generalizată",               "Dimensiuni generale, materiale principale"],
            ["300", "Definit",     "Specifică, cotată",          "Dimensiuni exacte, specificații materiale"],
            ["350", "Coordonare",  "LOD 300 + interfețe",        "Toate interfețele cu alte disciplini definite"],
            ["400", "Fabricare",   "Exactă (As-Built)",          "Toate detaliile, starea reală din teren"],
        ]
    )

    ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    safe = re.sub(r"[^\w]", "_", project)[:30]
    out  = GENERATED_DIR / f"LOD_{safe}_{ts}.docx"
    doc.save(str(out))
    return str(out)


def gen_eir(project: str) -> str:
    """Generează Employer's Information Requirements. Returnează calea fișierului."""
    ctx = get_project_context(project, [
        "beneficiar cerinte informatii proiect",
        "obiective BIM client investitor",
        "standarde ISO 19650 EIR cerinte",
        "livrabile predare documentatie",
    ])

    prompt = f"""Genereaza un EIR (Employer's Information Requirements) complet pentru proiectul: **{project}**

Context:
{ctx or "Proiect de constructii – EIR conform ISO 19650-2."}

## 1. Scopul EIR
## 2. Obiectivele BIM ale Beneficiarului
## 3. Cerinte de informatii organizationale (OIR)
## 4. Cerinte de informatii pentru activul construit (AIR)
## 5. Cerinte tehnice (software, formate, standarde)
## 6. Cerinte de management (CDE, procese, echipa)
## 7. Cerinte comerciale (termene, penalitati, livrabile)
## 8. Criterii de acceptanta
## 9. Anexe – Template-uri solicitate

Conform ISO 19650-2. Profesional, detaliat."""

    content = ask_claude(SYSTEM_BIM, prompt, max_tokens=3500)

    doc = _setup_doc("CERINȚE DE INFORMAȚII ALE BENEFICIARULUI (EIR)", project)
    _md_to_doc(doc, content)

    ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    safe = re.sub(r"[^\w]", "_", project)[:30]
    out  = GENERATED_DIR / f"EIR_{safe}_{ts}.docx"
    doc.save(str(out))
    return str(out)


def gen_requirements(project: str) -> str:
    """Extrage cerințe BIM din documentele proiectului. Returnează calea fișierului."""
    queries_map = {
        "Cerinte BIM generale":     "cerinte BIM modelare digitala proiect",
        "Software si platforme":    "software BIM Autodesk Revit IFC formate fisiere",
        "Livrabile BIM":            "livrabile BIM modele faze proiect predare",
        "Echipa BIM":               "BIM manager coordinator responsabilitati echipa",
        "CDE":                      "Common Data Environment CDE management informatii",
        "LOD":                      "LOD nivel detaliu modele BIM",
        "As-Built":                 "As-Built documentatie finala predare beneficiar",
        "Standarde":                "standarde BIM ISO 19650 aplicabile proiect",
        "Coordonare":               "coordonare interferente clash detection modele",
        "Obiectul contractului":    "obiectul contractului lucrari constructii",
    }

    doc = _setup_doc(f"EXTRAGERE CERINȚE BIM", project)
    _add_body(doc,
        f"Cerințe BIM extrase automat din documentele proiectului '{project}' "
        f"folosind căutare semantică (RAG) în baza de cunoștințe BIM."
    )

    found_any = False
    all_frags = []

    for topic, query in queries_map.items():
        full_q = f"{query} {project}" if project else query
        r = _query_rag(full_q, n=5)
        if not r.get("rag_used") or not r.get("context"):
            continue

        docs_list  = []
        metas_list = []
        for src in r.get("sources", []):
            if project.lower() in src.get("source", "").lower() or not project:
                metas_list.append(src)

        if not metas_list:
            metas_list = r.get("sources", [])

        if metas_list:
            found_any = True
            _add_h2(doc, topic)
            for src in metas_list[:3]:
                title = src.get("title", src.get("source", ""))
                page  = src.get("page", "-")
                rel   = src.get("relevance", 0)
                _add_body(doc, f"📄 {title} | Pag. {page} | Relevanță: {rel:.2f}")

            frags = r["context"].split("---")
            for frag in frags[:3]:
                frag = frag.strip()
                if len(frag) > 50:
                    _add_body(doc, frag[:600] + ("…" if len(frag) > 600 else ""))
                    all_frags.append(frag[:400])

    if not found_any:
        _add_body(doc,
            "Nu s-au găsit documente specifice acestui proiect în baza de date. "
            "Verificați că documentele au fost indexate cu bim_ingest.py."
        )
    else:
        # Rezumat generat de AI
        doc.add_page_break()
        _add_h1(doc, "Analiză și Concluzii")
        combined = "\n\n".join(all_frags[:15])
        summary = ask_claude(
            "Ești expert BIM în România. Analizezi documente de proiect.",
            f"Pe baza fragmentelor de mai jos din documentele proiectului '{project}', "
            f"rezumă cerințele BIM identificate și recomandă ce lipsește:\n\n{combined}",
            max_tokens=1500,
        )
        _md_to_doc(doc, summary)

    ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    safe = re.sub(r"[^\w]", "_", project)[:30]
    out  = GENERATED_DIR / f"CERINTE_{safe}_{ts}.docx"
    doc.save(str(out))
    return str(out)


def gen_checklist(project: str) -> str:
    """Generează Checklist Coordonare BIM. Returnează calea fișierului."""
    ctx = get_project_context(project, [
        "coordonare BIM clash detection interferente",
        "verificare calitate modele BIM",
        "sedinte BIM coordonare",
    ])

    prompt = f"""Genereaza un checklist detaliat de Coordonare BIM pentru proiectul: **{project}**

Context:
{ctx or "Proiect de constructii – checklist standard de coordonare BIM."}

## 1. Checklist Pre-Modelare (inainte de start)
## 2. Checklist Calitate Model (per disciplina)
## 3. Checklist Coordonare (federare modele, clash detection)
## 4. Checklist Sedinta BIM (agenda, minute)
## 5. Checklist Publicare in CDE
## 6. Checklist Receptie Model As-Built

Formatul fiecarui item: - [ ] Descriere actiune (Responsabil: X | Termen: Y)
Minim 50 de itemi total."""

    content = ask_claude(SYSTEM_BIM, prompt, max_tokens=3000)

    doc = _setup_doc("CHECKLIST COORDONARE BIM", project)
    _md_to_doc(doc, content)

    ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    safe = re.sub(r"[^\w]", "_", project)[:30]
    out  = GENERATED_DIR / f"CHECKLIST_{safe}_{ts}.docx"
    doc.save(str(out))
    return str(out)


def gen_minutes(project: str) -> str:
    """Generează template Minuta Ședință BIM. Returnează calea fișierului."""
    today = datetime.date.today().strftime("%d.%m.%Y")
    doc = _setup_doc("MINUTA ȘEDINȚĂ BIM", project)

    _add_h1(doc, "Date ședință")
    _add_table(doc, ["Câmp", "Valoare"], [
        ["Proiect",       project],
        ["Data",          today],
        ["Ora",           "________"],
        ["Locație / Link","________"],
        ["Moderator",     "BIM Manager – ________________"],
        ["Secretar",      "________________"],
        ["Nr. ședință",   "BIM-MTG-____"],
    ])

    _add_h1(doc, "Participanți")
    _add_table(doc,
        ["Nume", "Organizație", "Rol BIM", "Prezent"],
        [
            ["________________", "________________", "BIM Manager",      "□ Da  □ Nu"],
            ["________________", "________________", "BIM Coordinator",  "□ Da  □ Nu"],
            ["________________", "________________", "BIM Author Civil",  "□ Da  □ Nu"],
            ["________________", "________________", "Contractor BIM",    "□ Da  □ Nu"],
            ["________________", "________________", "Reprezentant Client","□ Da  □ Nu"],
            ["________________", "________________", "________________",  "□ Da  □ Nu"],
        ]
    )

    _add_h1(doc, "Agenda")
    _add_table(doc,
        ["Nr.", "Subiect", "Responsabil", "Durată"],
        [
            ["1", "Deschidere – prezenta si aprobarea agendei", "Moderator", "5 min"],
            ["2", "Revizie actiuni din sedinta anterioara",      "Moderator", "10 min"],
            ["3", "Status modele BIM – prezentare disciplini",   "BIM Coordinators", "15 min"],
            ["4", "Clash detection – probleme noi si rezolvate", "BIM Manager", "20 min"],
            ["5", "Coordonare RFI-uri deschise",                 "Contractor BIM", "10 min"],
            ["6", "Diverse",                                     "Toti", "5 min"],
            ["7", "Actiuni urmatoare si data urmatoarei sedinte","Moderator", "5 min"],
        ]
    )

    _add_h1(doc, "Discuții și Decizii")
    for i in range(1, 8):
        _add_h2(doc, f"Punctul {i} agenda")
        _add_body(doc, "________________________________________________________________________")
        _add_body(doc, "________________________________________________________________________")

    _add_h1(doc, "Plan de Acțiuni")
    _add_table(doc,
        ["#", "Acțiune", "Responsabil", "Termen", "Status"],
        [
            [str(i), "________________________________", "__________", "__________", "□ Deschis"]
            for i in range(1, 8)
        ]
    )

    _add_h1(doc, "Aprobare Minută")
    _add_table(doc, ["Rol", "Nume", "Semnătură", "Data"], [
        ["Moderator / BIM Manager", "________________", "________________", "________________"],
        ["Reprezentant Client",     "________________", "________________", "________________"],
    ])

    ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    safe = re.sub(r"[^\w]", "_", project)[:30]
    out  = GENERATED_DIR / f"MINUTA_{safe}_{ts}.docx"
    doc.save(str(out))
    return str(out)


def gen_iso_analysis(project: str) -> str:
    """Analiză conformitate ISO 19650. Returnează calea fișierului."""
    ctx = get_project_context(project, [
        "ISO 19650 cerinte standarde BIM",
        "OIR AIR PIR cerinte informatii",
        "CDE BEP EIR procese BIM",
    ])

    prompt = f"""Analizeaza conformitatea proiectului **{project}** cu SR EN ISO 19650.

Context din documentele proiectului:
{ctx or "Nu exista context specific."}

## 1. Rezumat conformitate (scor estimat pe 10 puncte)
## 2. ISO 19650-1 – Concepte si principii – Status
## 3. ISO 19650-2 – Faza de livrare – Status
   ### 3.1 Cerinte EIR – conformitate
   ### 3.2 BEP – conformitate
   ### 3.3 CDE – conformitate
   ### 3.4 Livrabile – conformitate
## 4. Lacune identificate (ce lipseste)
## 5. Recomandari prioritare (top 5 actiuni)
## 6. Roadmap conformitate ISO 19650

Fii critic si constructiv. Identifica lacunele reale."""

    content = ask_claude(SYSTEM_BIM, prompt, max_tokens=3000)

    doc = _setup_doc("ANALIZĂ CONFORMITATE ISO 19650", project)
    _md_to_doc(doc, content)

    ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    safe = re.sub(r"[^\w]", "_", project)[:30]
    out  = GENERATED_DIR / f"ISO19650_{safe}_{ts}.docx"
    doc.save(str(out))
    return str(out)


# ══════════════════════════════════════════════════════════════════════════════
# BEP PARAMETRIC — 13 capitole ISO 19650-2
# ══════════════════════════════════════════════════════════════════════════════

SYSTEM_BEP_PARAMETRIC = (
    "Ești BIM Manager senior în România, expert ISO 19650. "
    "Generezi secțiuni de BEP profesionale, adaptate la parametrii primiți. "
    "Răspunzi STRICT în JSON valid (fără ```json, fără text înainte/după). "
    "Limba: română. Fii detaliat dar concis."
)


def _ask_claude_json(system: str, user: str, max_tokens: int = 4000) -> dict:
    """Apel Claude cu parsare JSON. Fallback pe dict gol dacă parsing eșuează."""
    raw = ask_claude(system, user, max_tokens=max_tokens)
    # Strip markdown code fences if present
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        lines = cleaned.split("\n")
        lines = lines[1:]  # remove ```json
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        cleaned = "\n".join(lines)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return {"_raw_markdown": raw}


def gen_bep_parametric(params: dict) -> str:
    """Generează BEP parametric cu 13 capitole ISO 19650-2. Returnează calea fișierului."""
    import logging
    logger = logging.getLogger(__name__)

    project     = params.get("project", "Proiect BIM")
    client      = params.get("client", "De completat")
    work_type   = params.get("work_type", "constructie")
    phase       = params.get("phase", "PT")
    disciplines = params.get("disciplines", ["Arhitectura", "Structuri"])
    contractor  = params.get("contractor", "De desemnat")
    cde         = params.get("cde_platform", "ACC")
    standards   = params.get("standards", "SR EN ISO 19650-1/2, RTC 8")
    revit_data  = params.get("revit_data", "")
    today       = datetime.date.today().strftime("%d.%m.%Y")

    disc_str = ", ".join(disciplines)

    # ── RAG context ────────────────────────────────────────────────────────
    rag_queries = [
        f"BEP plan executie BIM {work_type}",
        f"cerinte BIM {disc_str}",
        f"standarde ISO 19650 {standards}",
        f"CDE {cde} structura foldere",
        f"LOD nivel detaliu {phase}",
    ]
    ctx = get_project_context(project, rag_queries, n=4)
    ctx_block = f"\n\nCONTEXT DIN DOCUMENTE:\n{ctx}" if ctx else ""

    # ── Common params for prompts ──────────────────────────────────────────
    params_block = (
        f"Proiect: {project}\n"
        f"Client/Beneficiar: {client}\n"
        f"Tip lucrare: {work_type}\n"
        f"Faza curentă: {phase}\n"
        f"Discipline: {disc_str}\n"
        f"Contractor execuție: {contractor}\n"
        f"Platformă CDE: {cde}\n"
        f"Standarde obligatorii: {standards}\n"
        f"Date suplimentare Revit: {revit_data or 'N/A'}\n"
        f"Data BEP: {today}"
    )

    # ── CALL 1: Capitolele 1-3 ─────────────────────────────────────────────
    logger.info("BEP parametric: Call 1/4 (Cap. 1-3)")
    c1 = _ask_claude_json(SYSTEM_BEP_PARAMETRIC, f"""Parametri proiect:
{params_block}
{ctx_block}

Generează JSON cu structura:
{{
  "info_rows": [["Câmp", "Valoare"], ...],  // 8-10 rânduri: beneficiar, proiectant, tip, fază, amplasament, nr.proiect, data BEP, versiune, status
  "scope_text": "Paragraf scop BEP...",
  "scope_bullets": ["bullet1", "bullet2", ...],
  "oir": ["obiectiv1", "obiectiv2", ...],  // 5-7 OIR
  "air": ["cerință1", "cerință2", ...],  // 4-6 AIR
  "pir_table": [["Nr.", "Cerință", "Faza", "Format", "Sursa"], ...],  // 8-12 rânduri
  "raci_table": [["Rol BIM", "Titular", "Responsabilitate", "RACI"], ...],  // 6-8 rânduri
  "bim_manager_duties": ["sarcina1", ...],  // 5-7 bullets
  "contractor_duties": ["sarcina1", ...]  // 5-6 bullets
}}

Adaptează conținutul la tipul de lucrare "{work_type}" și disciplinele "{disc_str}". Clientul este "{client}".""",
        max_tokens=4000)

    # ── CALL 2: Capitolele 4-6 ─────────────────────────────────────────────
    logger.info("BEP parametric: Call 2/4 (Cap. 4-6)")
    c2 = _ask_claude_json(SYSTEM_BEP_PARAMETRIC, f"""Parametri proiect:
{params_block}
{ctx_block}

Generează JSON cu structura:
{{
  "cde_info": [["Câmp", "Valoare"], ...],  // 6-8 rânduri: platformă, module, acces, admin, formate, backup
  "cde_folders": [["Folder CDE", "Conținut"], ...],  // 8-12 rânduri adaptate la discipline {disc_str}
  "cde_lifecycle": ["WIP description", "Shared description", "Published description", "Archived description"],
  "naming_convention": "Convenție denumire documente...",
  "standards_table": [["Standard", "Titlu", "Aplicabilitate"], ...],  // 8-10 rânduri
  "file_formats": [["Tip", "Format", "Software", "Utilizare"], ...],  // 7-9 rânduri
  "georef_text": "Paragraf georeferențiere cu EPSG:3844...",
  "software_table": [["Software", "Versiune min.", "Licență", "Utilizator", "Utilizare"], ...]  // 6-9 rânduri
}}

Platformă CDE: "{cde}". Discipline: "{disc_str}". Standarde obligatorii: "{standards}".""",
        max_tokens=4000)

    # ── CALL 3: Capitolele 7-9 ─────────────────────────────────────────────
    logger.info("BEP parametric: Call 3/4 (Cap. 7-9)")
    c3 = _ask_claude_json(SYSTEM_BEP_PARAMETRIC, f"""Parametri proiect:
{params_block}
{ctx_block}

Generează JSON cu structura:
{{
  "lod_intro": "Paragraf introductiv despre LOD...",
  "lod_definitions": [["LOD", "Denumire", "Descriere"], ...],  // LOD 100-400
  "lod_matrix": [["Element BIM", "PT (LOD)", "DDE (LOD)", "Execuție (LOD)", "As-Built (LOD)", "Note"], ...],  // 10-15 rânduri relevante pt {work_type}
  "deliverables_table": [["Cod", "Livrabil", "Format", "Responsabil", "Fază", "Destinatar"], ...],  // 10-14 rânduri BIM-LIV-XX
  "milestones_table": [["Jalon", "Eveniment", "Livrabile asociate"], ...],  // M0-M8
  "coordination_process": ["pas1", "pas2", ...],  // 5-6 pași flux coordonare
  "clash_types": [["Tip Clash", "Descriere", "Prioritate", "Termen remediere"], ...],  // 4 tipuri
  "meetings_text": ["frecventa_proiectare", "frecventa_executie", "milestone_review"]
}}

Discipline: "{disc_str}". Tip lucrare: "{work_type}". Faza curentă: "{phase}".""",
        max_tokens=4000)

    # ── CALL 4: Capitolele 10-13 ───────────────────────────────────────────
    logger.info("BEP parametric: Call 4/4 (Cap. 10-13)")
    c4 = _ask_claude_json(SYSTEM_BEP_PARAMETRIC, f"""Parametri proiect:
{params_block}
{ctx_block}

Generează JSON cu structura:
{{
  "asbuilt_requirements": ["cerința1", ...],  // 5-6 cerințe As-Built
  "technical_book": ["A. ...", "B. ...", ...],  // 4-5 componente carte tehnică
  "handover_items": ["item1", ...],  // 5-6 iteme predare
  "quality_checks": ["verificare1", ...],  // 5 verificări model
  "kpi_table": [["KPI", "Țintă", "Frecvență", "Responsabil"], ...],  // 5-6 KPI-uri
  "audit_text": ["audit_intern desc", "audit_extern desc"],
  "phase1_actions": ["acțiune1", ...],  // Faza 1 pregătire
  "phase2_actions": ["acțiune1", ...],  // Faza 2 proiectare
  "phase3_actions": ["acțiune1", ...],  // Faza 3 execuție
  "phase4_actions": ["acțiune1", ...],  // Faza 4 finalizare
  "approval_table": [["Rol", "Nume/Organizație", "Funcție", "Data", "Semnătură"], ...],  // 4-5 rânduri
  "revision_history": [["Versiune", "Data", "Modificări", "Aprobat de"], ...]  // 2-3 rânduri
}}

Client: "{client}". Contractor: "{contractor}". Tip: "{work_type}".""",
        max_tokens=3500)

    # ═══════════════════════════════════════════════════════════════════════
    # ASAMBLARE DOCX
    # ═══════════════════════════════════════════════════════════════════════
    logger.info("BEP parametric: Asamblare DOCX")
    doc = _setup_doc("PLAN DE EXECUȚIE BIM (BEP)", project)

    # Subtitlu
    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("BIM Execution Plan — ISO 19650-2")
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = GRAY_TEXT

    # Cover info table
    _add_info_table(doc, [
        ("Beneficiar / Client", client),
        ("Tip lucrare", work_type),
        ("Faza curentă", phase),
        ("Discipline", disc_str),
        ("Contractor execuție", contractor),
        ("Platformă CDE", cde),
        ("Data BEP", today),
        ("Versiune", "1.0 — Emitere Inițială"),
        ("Status", "DRAFT — pentru revizuire și aprobare"),
    ])

    p_std = doc.add_paragraph()
    r_std = p_std.add_run("Elaborat în conformitate cu SR EN ISO 19650-2 și bunele practici BIM")
    r_std.font.size = Pt(9)
    r_std.font.italic = True
    r_std.font.color.rgb = GRAY_TEXT

    doc.add_page_break()

    # ── Helper: fallback to markdown if JSON failed ────────────────────────
    def _has(d, key):
        return isinstance(d, dict) and key in d and d[key] and key != "_raw_markdown"

    def _fallback_md(d, doc_ref):
        if isinstance(d, dict) and "_raw_markdown" in d:
            _md_to_doc(doc_ref, d["_raw_markdown"])
            return True
        return False

    # ════════ CAPITOLUL 1: Informații Generale ══════════════════════════════
    _add_h1(doc, "1. INFORMAȚII GENERALE DESPRE PROIECT")

    if _has(c1, "info_rows"):
        _add_h2(doc, "1.1 Date de identificare")
        _add_info_table(doc, [(r[0], r[1]) for r in c1["info_rows"]])
    elif _fallback_md(c1, doc):
        pass
    else:
        _add_info_table(doc, [
            ("Denumire proiect", project),
            ("Beneficiar", client),
            ("Tip lucrare", work_type),
            ("Faza curentă", phase),
            ("Discipline", disc_str),
            ("Data BEP", today),
        ])

    if _has(c1, "scope_text"):
        _add_h2(doc, "1.2 Scopul BEP")
        _add_body(doc, c1["scope_text"])
        for b in c1.get("scope_bullets", []):
            _add_bullet(doc, b)

    doc.add_page_break()

    # ════════ CAPITOLUL 2: Obiective BIM ═══════════════════════════════════
    _add_h1(doc, "2. OBIECTIVELE BIM ȘI CERINȚELE DE INFORMAȚII")

    if _has(c1, "oir"):
        _add_h2(doc, "2.1 Obiective BIM ale Beneficiarului (OIR)")
        for o in c1["oir"]:
            _add_bullet(doc, o)

    if _has(c1, "air"):
        _add_h2(doc, "2.2 Cerințele de Informații ale Activului (AIR)")
        for a in c1["air"]:
            _add_bullet(doc, a)

    if _has(c1, "pir_table"):
        _add_h2(doc, "2.3 Cerințele de Informații ale Proiectului (PIR)")
        headers_pir = ["Nr.", "Cerință", "Faza", "Format", "Sursa"]
        _add_matrix_table(doc, headers_pir, c1["pir_table"])

    doc.add_page_break()

    # ════════ CAPITOLUL 3: Echipa BIM ══════════════════════════════════════
    _add_h1(doc, "3. ECHIPA BIM — ROLURI ȘI RESPONSABILITĂȚI")

    if _has(c1, "raci_table"):
        _add_h2(doc, "3.1 Matrice RACI — Echipa BIM")
        _add_matrix_table(doc, ["Rol BIM", "Titular / Organizație", "Responsabilitate principală", "RACI"],
                          c1["raci_table"])

    if _has(c1, "bim_manager_duties"):
        _add_h2(doc, "3.2 Responsabilitățile BIM Manager")
        for d in c1["bim_manager_duties"]:
            _add_bullet(doc, d)

    if _has(c1, "contractor_duties"):
        _add_h2(doc, "3.3 Responsabilitățile Contractor BIM")
        for d in c1["contractor_duties"]:
            _add_bullet(doc, d)

    doc.add_page_break()

    # ════════ CAPITOLUL 4: CDE ═════════════════════════════════════════════
    _add_h1(doc, "4. MEDIUL COMUN DE DATE (CDE)")

    if _has(c2, "cde_info"):
        _add_h2(doc, "4.1 Platforma CDE")
        _add_info_table(doc, [(r[0], r[1]) for r in c2["cde_info"]])
    elif _fallback_md(c2, doc):
        pass

    if _has(c2, "cde_folders"):
        _add_h2(doc, "4.2 Structura folderelor CDE")
        _add_matrix_table(doc, ["Folder CDE", "Conținut"], c2["cde_folders"])

    if _has(c2, "cde_lifecycle"):
        _add_h2(doc, "4.3 Ciclul de viață al documentelor")
        for status_desc in c2["cde_lifecycle"]:
            _add_bullet(doc, status_desc)

    if _has(c2, "naming_convention"):
        _add_body(doc, f"Convenție de denumire: {c2['naming_convention']}")

    doc.add_page_break()

    # ════════ CAPITOLUL 5: Standarde ═══════════════════════════════════════
    _add_h1(doc, "5. STANDARDE ȘI PROTOCOALE BIM")

    if _has(c2, "standards_table"):
        _add_h2(doc, "5.1 Standarde de referință")
        _add_matrix_table(doc, ["Standard", "Titlu", "Aplicabilitate"], c2["standards_table"])

    if _has(c2, "file_formats"):
        _add_h2(doc, "5.2 Formate de fișiere acceptate")
        _add_matrix_table(doc, ["Tip", "Format", "Software", "Utilizare"], c2["file_formats"])

    # Georeferentiere — deterministic section
    _add_h2(doc, "5.3 Sistemul de coordonate și georeferențiere")
    georef = c2.get("georef_text", "") if isinstance(c2, dict) else ""
    if georef:
        _add_body(doc, georef)
    else:
        _add_body(doc,
            "Toate modelele BIM vor fi georeferențiate în sistemul național de proiecție "
            "Stereografic 1970 (EPSG:3844), cu punct de origine comună stabilit de proiectantul "
            "topograf. Cota ±0.00 din modele va corespunde cotei absolute (mNMN)."
        )
    _add_body(doc, "Unitate de măsură: metru (m). Precizie coordonate: ±1 mm în model.")

    doc.add_page_break()

    # ════════ CAPITOLUL 6: Software ════════════════════════════════════════
    _add_h1(doc, "6. SOFTWARE ȘI PLATFORME")

    if _has(c2, "software_table"):
        _add_matrix_table(doc, ["Software / Platformă", "Versiune min.", "Licență", "Utilizator", "Utilizare"],
                          c2["software_table"])

    doc.add_page_break()

    # ════════ CAPITOLUL 7: LOD ═════════════════════════════════════════════
    _add_h1(doc, "7. NIVELELE DE DETALIU (LOD) ȘI INFORMAȚII (LOI)")

    if _has(c3, "lod_intro"):
        _add_body(doc, c3["lod_intro"])
    elif _fallback_md(c3, doc):
        pass

    if _has(c3, "lod_definitions"):
        _add_h2(doc, "7.1 Definiții LOD")
        _add_matrix_table(doc, ["LOD", "Denumire", "Descriere"], c3["lod_definitions"])

    if _has(c3, "lod_matrix"):
        _add_h2(doc, "7.2 Matrice LOD pe faze și elemente")
        _add_matrix_table(doc,
            ["Element BIM", "PT\n(LOD)", "DDE\n(LOD)", "Execuție\n(LOD)", "As-Built\n(LOD)", "Note"],
            c3["lod_matrix"])

    doc.add_page_break()

    # ════════ CAPITOLUL 8: Livrabile ═══════════════════════════════════════
    _add_h1(doc, "8. LIVRABILE BIM ȘI CALENDAR")

    if _has(c3, "deliverables_table"):
        _add_h2(doc, "8.1 Lista livrabilelor BIM")
        _add_matrix_table(doc,
            ["Cod", "Livrabil", "Format", "Responsabil", "Fază", "Destinatar"],
            c3["deliverables_table"])

    if _has(c3, "milestones_table"):
        _add_h2(doc, "8.2 Jaloane BIM (Milestones)")
        _add_matrix_table(doc,
            ["Jalon", "Eveniment", "Livrabile asociate"],
            c3["milestones_table"])

    doc.add_page_break()

    # ════════ CAPITOLUL 9: Coordonare ══════════════════════════════════════
    _add_h1(doc, "9. COORDONARE BIM ȘI CLASH DETECTION")

    if _has(c3, "coordination_process"):
        _add_h2(doc, "9.1 Procesul de coordonare")
        for i, step in enumerate(c3["coordination_process"], 1):
            _add_bullet(doc, f"{i}. {step}")

    if _has(c3, "clash_types"):
        _add_h2(doc, "9.2 Tipuri de interferențe verificate")
        _add_matrix_table(doc,
            ["Tip Clash", "Descriere", "Prioritate", "Termen remediere"],
            c3["clash_types"])

    if _has(c3, "meetings_text"):
        _add_h2(doc, "9.3 Ședințe BIM")
        for m in c3["meetings_text"]:
            _add_bullet(doc, m)

    doc.add_page_break()

    # ════════ CAPITOLUL 10: As-Built ═══════════════════════════════════════
    _add_h1(doc, "10. DOCUMENTAȚIE AS-BUILT ȘI PREDARE")

    if _has(c4, "asbuilt_requirements"):
        _add_h2(doc, "10.1 Cerințe As-Built")
        for r in c4["asbuilt_requirements"]:
            _add_bullet(doc, r)
    elif _fallback_md(c4, doc):
        pass

    if _has(c4, "technical_book"):
        _add_h2(doc, "10.2 Dosarul Cărții Tehnice (digital)")
        for t in c4["technical_book"]:
            _add_bullet(doc, t)

    if _has(c4, "handover_items"):
        _add_h2(doc, "10.3 Predare la beneficiar")
        for h in c4["handover_items"]:
            _add_bullet(doc, h)

    doc.add_page_break()

    # ════════ CAPITOLUL 11: Calitate ═══════════════════════════════════════
    _add_h1(doc, "11. MANAGEMENTUL CALITĂȚII BIM")

    if _has(c4, "quality_checks"):
        _add_h2(doc, "11.1 Verificări model (Quality Checks)")
        for q in c4["quality_checks"]:
            _add_bullet(doc, q)

    if _has(c4, "kpi_table"):
        _add_h2(doc, "11.2 Indicatori KPI BIM")
        _add_matrix_table(doc, ["KPI", "Țintă", "Frecvență", "Responsabil"], c4["kpi_table"])

    if _has(c4, "audit_text"):
        _add_h2(doc, "11.3 Audituri BIM")
        for a in c4["audit_text"]:
            _add_bullet(doc, a)

    doc.add_page_break()

    # ════════ CAPITOLUL 12: Plan implementare ══════════════════════════════
    _add_h1(doc, "12. PLAN DE IMPLEMENTARE BIM")

    for phase_num, (phase_title, key) in enumerate([
        ("Pregătire și configurare (M0)", "phase1_actions"),
        ("Proiectare PT/DDE (M0 → M2)", "phase2_actions"),
        ("Execuție (M3 → M6)", "phase3_actions"),
        ("Finalizare și predare (M7 → M8)", "phase4_actions"),
    ], 1):
        _add_h2(doc, f"12.{phase_num} Faza {phase_num} — {phase_title}")
        if _has(c4, key):
            for a in c4[key]:
                _add_bullet(doc, a)
        else:
            _add_body(doc, "De detaliat la kick-off meeting.")

    doc.add_page_break()

    # ════════ CAPITOLUL 13: Aprobări ═══════════════════════════════════════
    _add_h1(doc, "13. APROBĂRI BEP")

    _add_body(doc,
        "Prezentul BEP este supus aprobării tuturor părților implicate. "
        "Semnătura electronică sau umedă atestă acceptarea cerințelor.")

    if _has(c4, "approval_table"):
        _add_matrix_table(doc,
            ["Rol", "Nume / Organizație", "Funcție", "Data", "Semnătură"],
            c4["approval_table"])
    else:
        _add_matrix_table(doc,
            ["Rol", "Nume / Organizație", "Funcție", "Data", "Semnătură"],
            [
                ["Client / Beneficiar", client, "Reprezentant legal", "_________", "___________"],
                ["BIM Manager", "De desemnat", "BIM Manager", "_________", "___________"],
                ["Contractor BIM", contractor, "BIM Lead", "_________", "___________"],
            ])

    _add_h2(doc, "Istoricul reviziilor BEP")
    if _has(c4, "revision_history"):
        _add_matrix_table(doc,
            ["Versiune", "Data", "Modificări", "Aprobat de"],
            c4["revision_history"])
    else:
        _add_matrix_table(doc,
            ["Versiune", "Data", "Modificări", "Aprobat de"],
            [
                ["v1.0", today, "Emitere inițială — Draft", "BIM Manager"],
                ["v1.1", "TBD", "Actualizare după feedback", "BIM Manager + Client"],
            ])

    _add_body(doc,
        "Notă: BEP-ul va fi revizuit la fiecare jalon major (M0–M8) sau la modificări semnificative.")

    # ── Salvare ────────────────────────────────────────────────────────────
    ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    safe = re.sub(r"[^\w]", "_", project)[:30]
    out  = GENERATED_DIR / f"BEP_{safe}_{ts}.docx"
    doc.save(str(out))
    logger.info(f"BEP parametric generat: {out}")
    return str(out)


# ── Job management (pentru generare async) ────────────────────────────────────
_jobs: dict = {}
_jobs_lock = threading.Lock()

GENERATORS = {
    "bep":          gen_bep,
    "lod":          gen_lod,
    "eir":          gen_eir,
    "requirements": gen_requirements,
    "checklist":    gen_checklist,
    "minutes":      gen_minutes,
    "iso":          gen_iso_analysis,
}

DOC_LABELS = {
    "bep":          "Plan de Execuție BIM (BEP)",
    "lod":          "Matrice LOD / LOI",
    "eir":          "Cerințe Beneficiar (EIR)",
    "requirements": "Extragere Cerințe BIM",
    "checklist":    "Checklist Coordonare",
    "minutes":      "Minută Ședință BIM",
    "iso":          "Analiză Conformitate ISO 19650",
}


def start_generation(doc_type: str, project: str, params: dict = None) -> str:
    """Porneste generarea asincron. Returneaza job_id."""
    import uuid
    job_id = str(uuid.uuid4())[:8]
    with _jobs_lock:
        _jobs[job_id] = {"status": "running", "file": None, "error": None}

    def _run():
        try:
            if doc_type == "bep_parametric" and params:
                path = gen_bep_parametric(params)
            else:
                fn = GENERATORS[doc_type]
                path = fn(project)
            with _jobs_lock:
                _jobs[job_id]["status"] = "done"
                _jobs[job_id]["file"]   = Path(path).name
        except Exception as e:
            with _jobs_lock:
                _jobs[job_id]["status"] = "error"
                _jobs[job_id]["error"]  = str(e)

    threading.Thread(target=_run, daemon=True).start()
    return job_id


def get_job_status(job_id: str) -> dict:
    with _jobs_lock:
        return dict(_jobs.get(job_id, {"status": "not_found"}))


def list_generated_files() -> list:
    """Lista fisierelor generate, sortate dupa data (cele mai noi primele)."""
    files = []
    for f in sorted(GENERATED_DIR.glob("*.docx"), key=lambda x: x.stat().st_mtime, reverse=True):
        files.append({
            "name":     f.name,
            "size_kb":  round(f.stat().st_size / 1024, 1),
            "modified": datetime.datetime.fromtimestamp(f.stat().st_mtime).strftime("%d.%m.%Y %H:%M"),
        })
    return files

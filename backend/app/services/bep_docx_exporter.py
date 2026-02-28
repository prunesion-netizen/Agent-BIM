"""
bep_docx_exporter.py — Convertește Markdown BEP în document DOCX profesional.

Parsează Markdown linie cu linie și generează un DOCX formatat cu:
- Headings colorate (H1 albastru, H2 gri)
- Tabele cu header albastru și rânduri alternate
- Bullets, blockquotes, bold inline
- Cover page cu branding BIM
"""

import re
import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Culori ────────────────────────────────────────────────────────────────────
BLUE_DARK = RGBColor(0x1D, 0x4E, 0xD8)
BLUE_HEADER = "1D4ED8"
BLUE_ROW_EVEN = "F0F9FF"
WHITE_ROW_ODD = "FFFFFF"
GRAY_TEXT = RGBColor(0x37, 0x41, 0x51)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ── DOCX Helpers ──────────────────────────────────────────────────────────────
def _set_cell_shading(cell, hex_color: str):
    """Setează background color pe o celulă de tabel."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _setup_doc(project_code: str) -> Document:
    """Creează documentul DOCX cu margini și cover page."""
    doc = Document()

    # Margini profesionale
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.0)

    # Cover: linie decorativă
    p = doc.add_paragraph()
    r = p.add_run("\u2501" * 55)
    r.font.color.rgb = BLUE_DARK
    r.font.size = Pt(12)

    # Cover: titlu
    p2 = doc.add_paragraph()
    r2 = p2.add_run("BIM EXECUTION PLAN (BEP)")
    r2.font.size = Pt(20)
    r2.font.bold = True
    r2.font.color.rgb = BLUE_DARK

    # Cover: cod proiect
    p3 = doc.add_paragraph()
    r3 = p3.add_run(f"Proiect: {project_code}")
    r3.font.size = Pt(13)
    r3.font.color.rgb = GRAY_TEXT

    # Cover: data generării
    p4 = doc.add_paragraph()
    r4 = p4.add_run(
        f"Generat: {datetime.date.today().strftime('%d.%m.%Y')} \u00b7 Agent BIM Romania"
    )
    r4.font.size = Pt(9)
    r4.font.italic = True
    r4.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()  # spațiu
    return doc


def _add_heading(doc: Document, text: str, level: int):
    """Adaugă heading colorat."""
    p = doc.add_heading(text, level=level)
    if not p.runs:
        return
    run = p.runs[0]
    if level == 1:
        run.font.color.rgb = BLUE_DARK
        run.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.color.rgb = GRAY_TEXT
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
    elif level == 4:
        run.font.size = Pt(10)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)


def _add_body_paragraph(doc: Document, text: str):
    """Adaugă paragraf body cu formatare inline (bold)."""
    p = doc.add_paragraph()
    _add_inline_runs(p, text, font_size=Pt(10), font_color=GRAY_TEXT)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_bullet(doc: Document, text: str):
    """Adaugă bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    _add_inline_runs(p, text, font_size=Pt(10), font_color=GRAY_TEXT)
    return p


def _add_blockquote(doc: Document, text: str):
    """Adaugă blockquote ca paragraf italic cu indent."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    _add_inline_runs(p, text, font_size=Pt(10), font_color=GRAY_TEXT, italic=True)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_inline_runs(paragraph, text: str, font_size=None, font_color=None, italic=False):
    """Parsează **bold** inline și adaugă runs cu formatare."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.font.bold = True
        else:
            r = paragraph.add_run(part)
        if font_size:
            r.font.size = font_size
        if font_color:
            r.font.color.rgb = font_color
        if italic:
            r.font.italic = True


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Adaugă tabel cu header albastru și rânduri alternate."""
    if not headers:
        return
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"

    # Header row
    hr = t.rows[0]
    for j, h in enumerate(headers):
        _set_cell_shading(hr.cells[j], BLUE_HEADER)
        p = hr.cells[j].paragraphs[0]
        r = p.add_run(str(h).strip())
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, row in enumerate(rows):
        dr = t.rows[i + 1]
        bg = BLUE_ROW_EVEN if i % 2 == 0 else WHITE_ROW_ODD
        for j in range(len(headers)):
            cell_text = row[j] if j < len(row) else ""
            _set_cell_shading(dr.cells[j], bg)
            p = dr.cells[j].paragraphs[0]
            r = p.add_run(str(cell_text).strip())
            r.font.size = Pt(8.5)
            r.font.color.rgb = GRAY_TEXT
            if j == 0:
                r.font.bold = True

    doc.add_paragraph()  # spațiu după tabel


# ── Regex patterns ────────────────────────────────────────────────────────────
RE_H1 = re.compile(r"^# (.+)$")
RE_H2 = re.compile(r"^## (.+)$")
RE_H3 = re.compile(r"^### (.+)$")
RE_H4 = re.compile(r"^#{4,} (.+)$")
RE_BULLET = re.compile(r"^[\-\*] (.+)$")
RE_BLOCKQUOTE = re.compile(r"^> (.+)$")
RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
RE_TABLE_SEP = re.compile(r"^\|[\s\-:|]+\|$")
RE_HR = re.compile(r"^---+$")


def _parse_table_row(line: str) -> list[str]:
    """Extrage celulele dintr-un rând de tabel Markdown."""
    return [c.strip() for c in line.strip("|").split("|")]


# ── Funcția principală ────────────────────────────────────────────────────────
def markdown_to_docx(markdown: str, project_code: str) -> BytesIO:
    """
    Convertește un string Markdown BEP într-un document DOCX profesional.

    Returns:
        BytesIO buffer cu fișierul DOCX gata de descărcat.
    """
    doc = _setup_doc(project_code)
    lines = markdown.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip linii goale
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if RE_HR.match(line.strip()):
            i += 1
            continue

        # Headings (de la H4 în sus, verificăm cel mai specific mai întâi)
        m = RE_H4.match(line.strip())
        if m:
            _add_heading(doc, m.group(1), 4)
            i += 1
            continue

        m = RE_H3.match(line.strip())
        if m:
            _add_heading(doc, m.group(1), 3)
            i += 1
            continue

        m = RE_H2.match(line.strip())
        if m:
            _add_heading(doc, m.group(1), 2)
            i += 1
            continue

        m = RE_H1.match(line.strip())
        if m:
            _add_heading(doc, m.group(1), 1)
            i += 1
            continue

        # Tabel Markdown — colectează toate rândurile tabelului
        if RE_TABLE_ROW.match(line.strip()):
            table_lines = []
            while i < len(lines) and RE_TABLE_ROW.match(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 2:
                headers = _parse_table_row(table_lines[0])
                data_rows = []
                for tl in table_lines[1:]:
                    if not RE_TABLE_SEP.match(tl):
                        data_rows.append(_parse_table_row(tl))
                _add_table(doc, headers, data_rows)
            continue

        # Bullet
        m = RE_BULLET.match(line.strip())
        if m:
            _add_bullet(doc, m.group(1))
            i += 1
            continue

        # Blockquote
        m = RE_BLOCKQUOTE.match(line.strip())
        if m:
            _add_blockquote(doc, m.group(1))
            i += 1
            continue

        # Text normal
        _add_body_paragraph(doc, line.strip())
        i += 1

    # Footer cu data generării
    doc.add_paragraph()
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run(
        f"\u2500\u2500\u2500 Document generat automat \u00b7 "
        f"{datetime.date.today().strftime('%d.%m.%Y')} \u00b7 "
        f"Agent BIM Romania \u2500\u2500\u2500"
    )
    rf.font.size = Pt(8)
    rf.font.italic = True
    rf.font.color.rgb = GRAY_TEXT

    # Salvează în buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

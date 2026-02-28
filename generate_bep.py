"""
generate_bep.py — Generează BEP (BIM Execution Plan) pentru proiectul Ghizela Celula 3
Produce: BEP_Ghizela_Celula3.docx

Rulare: python generate_bep.py
"""
import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT_FILE = "BEP_Ghizela_Celula3.docx"
TODAY = datetime.date.today().strftime("%d.%m.%Y")

# ── Culori brand BIM ────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1D, 0x4E, 0xD8)   # #1D4ED8 albastru accent
BLUE_LIGHT = RGBColor(0xDB, 0xEA, 0xFE)   # fundal header tabel
GRAY_TEXT  = RGBColor(0x37, 0x41, 0x51)   # gri text body
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_background(cell, hex_color: str):
    """Setează culoarea de fundal a celulei."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Adaugă chenar la o celulă."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), kwargs.get("val", "single"))
        border.set(qn("w:sz"), str(kwargs.get("sz", 4)))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), kwargs.get("color", "1D4ED8"))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1, color=None):
    """Adaugă heading cu stil personalizat."""
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if color:
        run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, bold=False, italic=False, size=10):
    """Adaugă paragraf body."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = GRAY_TEXT
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0):
    """Adaugă bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_TEXT
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p


def add_info_table(doc, rows: list):
    """Tabel cheie-valoare pentru informații proiect."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [Cm(5.5), Cm(11.5)]
    for i, (key, val) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]

        # Cheie (albastru deschis)
        set_cell_background(row.cells[0], "DBF0FE")
        kp = row.cells[0].paragraphs[0]
        kr = kp.add_run(key)
        kr.font.bold = True
        kr.font.size = Pt(9)
        kr.font.color.rgb = BLUE_DARK

        # Valoare
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(val)
        vr.font.size = Pt(9)
        vr.font.color.rgb = GRAY_TEXT
    doc.add_paragraph()


def add_matrix_table(doc, headers: list, data: list):
    """Tabel matrice cu header albastru."""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = "Table Grid"

    # Header
    hrow = table.rows[0]
    for j, h in enumerate(headers):
        set_cell_background(hrow.cells[j], "1D4ED8")
        p = hrow.cells[j].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    for i, row_data in enumerate(data):
        drow = table.rows[i + 1]
        bg = "F0F9FF" if i % 2 == 0 else "FFFFFF"
        for j, cell_text in enumerate(row_data):
            set_cell_background(drow.cells[j], bg)
            p = drow.cells[j].paragraphs[0]
            r = p.add_run(str(cell_text))
            r.font.size = Pt(8.5)
            r.font.color.rgb = GRAY_TEXT
            if j == 0:
                r.font.bold = True
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════════
def main():
    doc = Document()

    # ── Marjă pagină ─────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    # ════════════════════════════════════════════════════════════════════════════
    # COPERTĂ
    # ════════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    # Linie albastră sus
    cover_line = doc.add_paragraph()
    cover_line_run = cover_line.add_run("━" * 55)
    cover_line_run.font.color.rgb = BLUE_DARK
    cover_line_run.font.size = Pt(14)

    p_type = doc.add_paragraph()
    r = p_type.add_run("PLAN DE EXECUȚIE BIM")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = BLUE_DARK
    p_type.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_sub = doc.add_paragraph()
    r2 = p_sub.add_run("BIM Execution Plan (BEP)")
    r2.font.size = Pt(14)
    r2.font.italic = True
    r2.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()

    p_proj = doc.add_paragraph()
    r3 = p_proj.add_run("Construire Celula 3 din Depozitul de deșeuri\nnepericuloase Ghizela")
    r3.font.size = Pt(17)
    r3.font.bold = True
    r3.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()
    doc.add_paragraph()

    cover_meta = [
        ("Beneficiar",       "Unitatea Administrativ Teritorială – Județul Timiș (CJ Timiș)"),
        ("Proiectant General", "COPLAN CAI SRL"),
        ("Nr. Proiect",      "426 / 03.07.2023"),
        ("Nr. Contract",     "147 / 17.12.2025"),
        ("Amplasament",      "Comuna Ghizela, Județul Timiș"),
        ("Data BEP",         TODAY),
        ("Versiune",         "1.0 – Emitere Inițială"),
        ("Status",           "DRAFT – pentru revizuire și aprobare"),
    ]
    add_info_table(doc, cover_meta)

    p_std = doc.add_paragraph()
    r4 = p_std.add_run("Elaborat în conformitate cu SR EN ISO 19650-2:2021 și bunele practici BIM")
    r4.font.size = Pt(9)
    r4.font.italic = True
    r4.font.color.rgb = GRAY_TEXT

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 1. INFORMAȚII GENERALE DESPRE PROIECT
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. INFORMAȚII GENERALE DESPRE PROIECT", 1, BLUE_DARK)

    add_heading(doc, "1.1 Date de identificare", 2)
    add_info_table(doc, [
        ("Denumire proiect",   "Construire Celula 3 din Depozitul de deșeuri nepericuloase Ghizela"),
        ("Tip lucrare",        "Construire celulă depozit deșeuri nepericuloase (infrastructură mediu)"),
        ("Amplasament",        "Comuna Ghizela, Județul Timiș, România"),
        ("Beneficiar (Client)", "Unitatea Administrativ Teritorială – Județul Timiș, prin Consiliul Județean Timiș"),
        ("Proiectant General", "COPLAN CAI SRL"),
        ("Nr. Proiect",        "426 / 03.07.2023"),
        ("Nr. Contract execuție", "147 / 17.12.2025"),
        ("Data Proiect Tehnic", "21.02.2025"),
        ("Faza proiect curentă", "PT + CS + DE (Proiect Tehnic, Caiet de Sarcini, Detalii de Execuție)"),
    ])

    add_heading(doc, "1.2 Parametri tehnici principali", 2)
    add_info_table(doc, [
        ("Suprafață celulă (la bază)", "~7 ha"),
        ("Înălțime maximă umplere",    "20 m"),
        ("Capacitate proiectată",       "Celula 3 – extensie a Celulei 2 existente"),
        ("Tip deșeuri",                "Deșeuri nepericuloase"),
        ("Reglementare principală",    "HG 349/2005, Ordinul MMGA 757/2004, OUG 195/2005, Acord mediu 6/18.09.2009"),
        ("Acord de mediu",             "Nr. 6/18.09.2009, revizuit 26.08.2010 și 19.04.2011"),
    ])

    add_heading(doc, "1.3 Scopul BEP", 2)
    add_body(doc,
        "Prezentul Plan de Execuție BIM (BEP) definește modul în care echipa de proiect va "
        "implementa cerințele privind managementul informațiilor pe parcursul proiectului "
        "\"Construire Celula 3 din Depozitul de deșeuri nepericuloase Ghizela\", în conformitate cu "
        "SR EN ISO 19650-2:2021."
    )
    add_body(doc, "BEP-ul stabilește:")
    add_bullet(doc, "Structura și responsabilitățile echipei BIM")
    add_bullet(doc, "Standardele, protocoalele și formatele de livrare a informațiilor")
    add_bullet(doc, "Mediul Comun de Date (CDE) utilizat și regulile de gestionare")
    add_bullet(doc, "Nivelele de Detaliu (LOD) și de Informații (LOI) solicitate pe faze")
    add_bullet(doc, "Livrabilele BIM și calendarul de predare")
    add_bullet(doc, "Procesul de coordonare și detectare interferențe (clash detection)")
    add_bullet(doc, "Cerințele pentru documentația As-Built și predarea la beneficiar")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 2. OBIECTIVELE BIM ȘI CERINȚELE DE INFORMAȚII
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. OBIECTIVELE BIM ȘI CERINȚELE DE INFORMAȚII", 1, BLUE_DARK)

    add_heading(doc, "2.1 Obiective BIM ale Beneficiarului (OIR)", 2)
    add_body(doc,
        "Consiliul Județean Timiș, în calitate de client, urmărește prin implementarea BIM "
        "atingerea următoarelor obiective organizaționale (Organizational Information Requirements – OIR):"
    )
    add_bullet(doc, "Vizualizare 3D și înțelegere clară a soluțiilor tehnice propuse (etanșare, drenaj levigat, diguri)")
    add_bullet(doc, "Monitorizarea progresului execuției față de planificare (4D BIM)")
    add_bullet(doc, "Control cantitativ și estimare costuri îmbunătățite (5D BIM – Takeoff)")
    add_bullet(doc, "Documentație As-Built digitală pentru exploatare și monitorizare post-execuție")
    add_bullet(doc, "Arhivare structurată a documentației tehnice pe platformă CDE accesibilă pe termen lung")
    add_bullet(doc, "Conformitate cu cerințele de audit și raportare pentru fonduri europene (dacă aplicabil)")

    add_heading(doc, "2.2 Cerințele de Informații ale Activului (AIR)", 2)
    add_body(doc,
        "Cerințele de informații pentru activul construit (Asset Information Requirements – AIR) "
        "vizează suportul operațional post-execuție al depozitului:"
    )
    add_bullet(doc, "Model geometric 3D al celulei finalizate (straturi etanșare, drenaj, diguri, rampe acces)")
    add_bullet(doc, "Baza de date a materialelor geosintétice (producător, specificații, certificate conformitate)")
    add_bullet(doc, "Planuri As-Built georeferențiate pentru monitorizarea tasărilor și subsidențelor")
    add_bullet(doc, "Documentație exploatare: rețea levigat, cămine CL2.15/CL1.15, conducte descărcare")
    add_bullet(doc, "Fișe tehnice echipamente monitorizare (senzori, piezometre, puțuri de monitoring)")

    add_heading(doc, "2.3 Cerințele de Informații ale Proiectului (PIR)", 2)
    add_body(doc,
        "Cerințele de informații specifice proiectului (Project Information Requirements – PIR), "
        "derivate din documentele contractuale (CS, PT, DTOE):"
    )
    add_matrix_table(doc,
        ["Nr.", "Cerință", "Faza", "Format", "Sursa"],
        [
            ["1", "Model 3D terasamente și profilare teren", "PT/DDE", "RVT / IFC", "CS Art. 2"],
            ["2", "Model 3D sistem etanșare (geomembrană HDPE, geotextile)", "DDE", "RVT / IFC", "CS Cap. Etanșare"],
            ["3", "Model 3D rețea drenaj levigat (conducte, cămine)", "DDE", "RVT / IFC", "Memoriu PT"],
            ["4", "Model 3D diguri perimetrale și rampe acces", "PT/DDE", "RVT / IFC", "DTOE"],
            ["5", "Cantități terasamente (săpătură, umplutură)", "PT", "Excel / XLSX", "CS Art. 3"],
            ["6", "Planuri coordonare multi-disciplinară", "DDE", "NWD / BCF", "ISO 19650-2"],
            ["7", "Raport clash detection (interferențe)", "DDE", "BCF / PDF", "ISO 19650-2"],
            ["8", "Model As-Built complet", "Recepție", "RVT / IFC", "CS Art. 16-17"],
            ["9", "Dosare de carte tehnică digitală", "Recepție", "PDF / DWF", "HG 273/94"],
            ["10", "Planuri de monitorizare post-închidere", "Recepție", "PDF / DWG", "HG 349/2005"],
        ]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 3. ECHIPA BIM – ROLURI ȘI RESPONSABILITĂȚI
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. ECHIPA BIM – ROLURI ȘI RESPONSABILITĂȚI", 1, BLUE_DARK)

    add_body(doc,
        "Structura echipei BIM pentru proiectul Ghizela Celula 3 este organizată conform "
        "SR EN ISO 19650-2:2021, Anexa B, cu responsabilități clare pe niveluri de decizie și execuție."
    )

    add_heading(doc, "3.1 Matrice RACI – Echipa BIM", 2)
    add_matrix_table(doc,
        ["Rol BIM", "Titular / Organizație", "Responsabilitate principală", "RACI"],
        [
            ["Client / Beneficiar", "CJ Timiș", "Aprobă BEP, validează livrabile finale, asigură accesul la CDE", "A"],
            ["BIM Manager Proiect", "COPLAN CAI SRL\n(desemnat)", "Coordonare generală BIM, menținere BEP, admin CDE, raportare", "R"],
            ["BIM Coordinator\nStructură", "COPLAN CAI SRL", "Model 3D terasamente, etanșare, diguri; clash detection structuri", "R"],
            ["BIM Coordinator\nInstalaţii/Drenaj", "Proiectant instalații\n(subcontractor)", "Model 3D rețea levigat, conducte, cămine; coordonare cu structură", "R"],
            ["BIM Author\nArchitect/Civil", "COPLAN CAI SRL", "Modelare obiecte Revit, respectare LOD, publicare IFC în CDE", "R"],
            ["Contractor BIM Lead", "Inoveco SRL\n(desemnat)", "Urmărire model în teren, actualizare model execuție, 4D planning", "R"],
            ["Diriginte de Șantier", "CJ Timiș / Desemnat", "Validare conformitate model vs. teren, semnare PV recepție", "C"],
            ["CDE Administrator", "COPLAN CAI SRL\n(sau Inoveco)", "Administrare permisiuni ACC, structură foldere, audit trail", "R"],
        ]
    )

    add_heading(doc, "3.2 Responsabilitățile BIM Manager", 2)
    add_bullet(doc, "Elaborarea, actualizarea și distribuirea BEP-ului pe parcursul proiectului")
    add_bullet(doc, "Stabilirea și menținerea structurii CDE (Autodesk ACC Docs)")
    add_bullet(doc, "Organizarea sesiunilor de coordonare BIM (săptămânale sau bi-săptămânale)")
    add_bullet(doc, "Verificarea calității modelelor înainte de publicarea în zona Shared a CDE")
    add_bullet(doc, "Gestionarea rapoartelor de clash detection și urmărirea remedierii interferențelor")
    add_bullet(doc, "Raportarea statusului BIM la beneficiar (lunar)")
    add_bullet(doc, "Asigurarea conformității cu SR EN ISO 19650 și cerințele contractuale")

    add_heading(doc, "3.3 Responsabilitățile Contractor BIM Lead (Inoveco)", 2)
    add_bullet(doc, "Utilizarea platformei ACC Build pentru urmărirea lucrărilor și RFI-uri")
    add_bullet(doc, "Actualizarea modelelor de execuție cu devieri față de proiect (as-built progresiv)")
    add_bullet(doc, "Conectarea activităților din program (planificare 4D) la elementele din model")
    add_bullet(doc, "Utilizarea ACC Takeoff pentru verificarea cantităților față de proiect")
    add_bullet(doc, "Documentarea foto și video a fazelor de lucrări ascunse în ACC Docs")
    add_bullet(doc, "Predarea modelului As-Built complet la recepția preliminară")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 4. MEDIUL COMUN DE DATE (CDE)
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. MEDIUL COMUN DE DATE (CDE)", 1, BLUE_DARK)

    add_heading(doc, "4.1 Platforma CDE", 2)
    add_info_table(doc, [
        ("Platformă CDE",       "Autodesk Construction Cloud (ACC)"),
        ("Module utilizate",    "ACC Docs | ACC Coordinate | ACC Build | ACC Takeoff"),
        ("Acces",               "Web browser (https://acc.autodesk.com) + aplicații mobile iOS/Android"),
        ("Administrator CDE",   "COPLAN CAI SRL / BIM Manager desemnat"),
        ("Formate suportate",   "RVT, IFC, DWG, DWF, PDF, XLSX, DOCX, JPG, BCF"),
        ("Audit trail",         "Automat – toate acțiunile sunt înregistrate cu utilizator și timestamp"),
        ("Backup",              "Autodesk cloud – redundanță geografică automată"),
    ])

    add_heading(doc, "4.2 Structura folderelor CDE", 2)
    add_body(doc, "Structura de foldere în ACC Docs va respecta următoarea ierarhie:")

    folders = [
        ("00_BEP și Standarde",        "BEP curent, protocoale, standarde de referință, template-uri"),
        ("01_Modele BIM",               "Subfoldere: /WIP | /Shared | /Published | /Archived"),
        ("01_Modele BIM / WIP",         "Modele în lucru (acces restricționat per disciplină)"),
        ("01_Modele BIM / Shared",      "Modele partajate pentru coordonare (după verificare BIM Manager)"),
        ("01_Modele BIM / Published",   "Versiuni aprobate și livrabile oficiale"),
        ("02_Documente Proiect",        "PT, CS, DE, avize, acorduri, contract, DTOE"),
        ("03_Coordonare BIM",           "Rapoarte clash detection, procese verbale coordonare, BCF"),
        ("04_Execuție (Build)",         "RFI-uri, submittal-uri, devize, situații de lucrări"),
        ("05_Cantități (Takeoff)",      "Modele Takeoff, extrase cantități, comparații"),
        ("06_As-Built",                 "Modele și planuri finale As-Built, fotografii execuție"),
        ("07_Recepție",                 "Dosare carte tehnică, procese verbale recepție, garanții"),
        ("08_Comunicări",               "Corespondență oficială, minute ședințe, notificări"),
    ]
    add_matrix_table(doc,
        ["Folder CDE", "Conținut"],
        [[f, d] for f, d in folders]
    )

    add_heading(doc, "4.3 Ciclul de viață al documentelor în CDE", 2)
    add_body(doc, "Documentele parcurg obligatoriu următoarele stări (status) în CDE:")
    add_bullet(doc, "WIP (Work In Progress) – document în lucru, vizibil doar echipei autorului")
    add_bullet(doc, "Shared – partajat pentru revizuire de către BIM Coordinator / BIM Manager")
    add_bullet(doc, "Published – aprobat și publicat oficial; nu poate fi modificat fără nouă versiune")
    add_bullet(doc, "Archived – versiune supercedată, păstrată pentru audit trail")
    add_body(doc, "\nConvenție de numerotare documente:", bold=True)
    add_body(doc, "[COD-PROIECT]-[DISCIPLINA]-[TIP]-[NUMĂR]-[REVIZIE]")
    add_body(doc, "Exemplu: GHZ-STR-DRW-0001-C01 (Ghizela – Structuri – Desen – nr. 1 – Revizie C01)")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 5. STANDARDE ȘI PROTOCOALE BIM
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. STANDARDE ȘI PROTOCOALE BIM", 1, BLUE_DARK)

    add_heading(doc, "5.1 Standarde de referință", 2)
    add_matrix_table(doc,
        ["Standard", "Titlu", "Aplicabilitate"],
        [
            ["SR EN ISO 19650-1:2019", "Organizarea și digitalizarea informațiilor despre clădiri și lucrări de inginerie civilă – Managementul informațiilor cu ajutorul BIM – Concepte și principii", "Cadru general BIM"],
            ["SR EN ISO 19650-2:2021", "Faza de livrare a activelor", "Cerințe BEP, CDE, livrabile"],
            ["SR EN ISO 19650-3:2021", "Faza operațională a activelor (AssetIM)", "As-Built și exploatare"],
            ["ISO 16739 (IFC 4.3)", "Industry Foundation Classes – format deschis de schimb", "Export/import modele"],
            ["BS EN 17412-1:2021", "Level of Information Need (LOiN)", "Definire LOD/LOI"],
            ["RTC 8 (2023)", "Referentul Tehnic Constructii – Categoria 8", "Verificare tehnică proiect"],
            ["RTC 9 (2023)", "Referentul Tehnic Constructii – Categoria 9", "Verificare structuri speciale"],
            ["HG 273/1994", "Regulament recepție lucrări de construcții", "Recepție preliminară și finală"],
            ["HG 349/2005", "Depozitarea deșeurilor – cerințe tehnice", "Cerințe specifice depozit"],
            ["Ord. MMGA 757/2004", "Normativ tehnic privind depozitarea deșeurilor", "Standarde construcție depozit"],
        ]
    )

    add_heading(doc, "5.2 Formate de fișiere acceptate", 2)
    add_matrix_table(doc,
        ["Tip", "Format", "Software", "Utilizare"],
        [
            ["Model nativ BIM",   "RVT",       "Autodesk Revit",         "Modelare, coordonare intern"],
            ["Model de schimb",   "IFC 2x3/4", "Orice software BIM",     "Interoperabilitate, predare client"],
            ["Model coordonare",  "NWD/NWC",   "Autodesk Navisworks",    "Clash detection, 4D"],
            ["Raport interferențe", "BCF 2.1", "Navisworks/BIM Collab",  "Comunicare probleme clash"],
            ["Planuri 2D",        "DWG / PDF", "AutoCAD / Revit",        "Planuri execuție, As-Built"],
            ["Cantități",         "XLSX",      "ACC Takeoff / Excel",    "Extrase cantități, devize"],
            ["Fotografii",        "JPG/PNG",   "ACC Build / mobil",      "Documentare teren"],
            ["Documente text",    "PDF / DOCX", "Office / Acrobat",      "Rapoarte, PV, certificate"],
            ["Georeferențiere",   "DWG / SHP", "AutoCAD Civil 3D / GIS", "Planuri topografice"],
        ]
    )

    add_heading(doc, "5.3 Sistemul de coordonate și georeferențiere", 2)
    add_body(doc,
        "Toate modelele BIM vor fi georeferențiate în sistemul național de proiecție "
        "Stereografic 1970 (EPSG:3844), cu punct de origine comună stabilit de proiectantul "
        "topograf. Cota ±0.00 din modele va corespunde cotei absolute (mNMN) precizată în "
        "planurile topografice."
    )
    add_body(doc,
        "Unitate de măsură: metru (m). Precizie coordonate: ±1 mm în model."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 6. SOFTWARE ȘI PLATFORME
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. SOFTWARE ȘI PLATFORME", 1, BLUE_DARK)

    add_matrix_table(doc,
        ["Software / Platformă", "Versiune minimă", "Licență", "Utilizator", "Utilizare"],
        [
            ["Autodesk Revit",         "2024",       "Autodesk AEC Collection", "Proiectant BIM",      "Modelare 3D arhitectură, structuri, instalații"],
            ["Autodesk AutoCAD Civil 3D", "2024",    "Autodesk AEC Collection", "Proiectant Civil",    "Modele teren, trasare, volume terasamente"],
            ["Autodesk Navisworks Manage", "2024",   "Autodesk AEC Collection", "BIM Coordinator",     "Clash detection, 4D scheduling, model federation"],
            ["ACC Docs",               "Web/Cloud",  "ACC License",             "Toate echipele",      "Document management, CDE principal"],
            ["ACC Coordinate",         "Web/Cloud",  "ACC License",             "BIM Manager/Coord.",  "Coordonare modele 3D, clash cloud"],
            ["ACC Build",              "Web/Mobile", "ACC License",             "Contractor Inoveco",  "RFI-uri, submittal, punch lists, inspecții"],
            ["ACC Takeoff",            "Web/Cloud",  "ACC License",             "Estimator/Contractor", "Extragere cantități din modele 2D/3D"],
            ["BIM Collab Zoom (opțional)", "Web",   "BIM Collab",              "BIM Manager",         "Coordonare BCF issues, sesiuni BIM"],
            ["Microsoft Project / Primavera", "-",   "Client / Contractor",     "Contractor",          "Planificare 4D, baseline program"],
        ]
    )

    add_heading(doc, "6.1 Gestionarea versiunilor software", 2)
    add_body(doc,
        "Toți membrii echipei BIM vor utiliza aceleași versiuni de software pentru a evita "
        "incompatibilitățile la schimbul de fișiere. Orice upgrade de versiune se comunică "
        "BIM Manager-ului cu minim 2 săptămâni înainte, iar schimbul de modele IFC este "
        "metoda de rezervă pentru echipele cu versiuni diferite."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 7. NIVELELE DE DETALIU (LOD) ȘI INFORMAȚII (LOI)
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. NIVELELE DE DETALIU (LOD) ȘI INFORMAȚII (LOI)", 1, BLUE_DARK)

    add_body(doc,
        "Nivelele de Detaliu Geometric (LOG) și de Informații (LOI) sunt definite conform "
        "BS EN 17412-1:2021 (Level of Information Need) și sunt adaptate fazelor proiectului "
        "Ghizela Celula 3. Convențional, se utilizează scala LOD 100–400 (RIBA/BIMForum)."
    )

    add_heading(doc, "7.1 Matrice LOD pe faze și elemente", 2)
    add_matrix_table(doc,
        ["Element BIM", "PT\n(LOD)", "DDE\n(LOD)", "Execuție\n(LOD)", "As-Built\n(LOD)", "Note"],
        [
            ["Corp depozit – contur celulă", "200", "300", "350", "400", "Coordonate reale din topo"],
            ["Sistem etanșare bază (HDPE + GCL)", "200", "350", "350", "400", "Inclusiv specificații material"],
            ["Geotextile protecție", "200", "300", "350", "400", "Producător, gramaj, suprafață"],
            ["Diguri perimetrale", "200", "300", "350", "400", "Profile transversale reale"],
            ["Drenaj levigat – conducte principale", "200", "300", "350", "400", "Diametre, materiale, pante"],
            ["Cămine levigat (CL)", "200", "300", "350", "400", "Coordonate, tip, dimensiuni"],
            ["Rampe acces subcelule", "200", "300", "350", "400", "Structură: geogrile, balast"],
            ["Sistem colectare ape meteorice", "100", "300", "350", "400", "Șanțuri, deversoare"],
            ["Echipamente monitorizare", "100", "200", "350", "400", "Senzori, piezometre"],
            ["Organizare șantier", "100", "300", "—", "—", "Conform DTOE"],
        ]
    )

    add_heading(doc, "7.2 Definiții LOD utilizate", 2)
    add_matrix_table(doc,
        ["LOD", "Denumire", "Descriere"],
        [
            ["100", "Concept",      "Element reprezentat simbolic; suprafețe, volume aproximative; fără geometrie exactă"],
            ["200", "Schematic",    "Element cu geometrie generalizată; dimensiuni, locație și orientare aproximative"],
            ["300", "Definit",      "Element cu geometrie specifică; dimensiuni, materiale și locație exacte; baza pentru fabricare"],
            ["350", "Coordonare",   "LOD 300 + interfețe cu alte elemente definite; gata pentru clash detection"],
            ["400", "Fabricare",    "Element complet cu toate detaliile de execuție; corespunde stării reale As-Built"],
        ]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 8. LIVRABILE BIM ȘI CALENDAR
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. LIVRABILE BIM ȘI CALENDAR", 1, BLUE_DARK)

    add_heading(doc, "8.1 Lista livrabilelor BIM", 2)
    add_matrix_table(doc,
        ["Cod", "Livrabil", "Format", "Responsabil", "Fază", "Destinatar"],
        [
            ["BIM-LIV-01", "BEP v1.0 – Emitere Inițială (prezentul document)", "DOCX/PDF", "BIM Manager", "Inițiere", "Client, Contractor"],
            ["BIM-LIV-02", "Model federat PT – Revit federat + IFC 4.3", "RVT + IFC", "BIM Manager", "PT finalizat", "Client, Contractor"],
            ["BIM-LIV-03", "Planuri 2D extrase din model (secțiuni, vederi)", "DWG + PDF", "BIM Author", "PT finalizat", "Execuție"],
            ["BIM-LIV-04", "Extrase cantități terasamente (din model Civil 3D)", "XLSX + PDF", "BIM Author Civil", "PT finalizat", "Contractor, Client"],
            ["BIM-LIV-05", "Model coordonare (Navisworks federat)", "NWD", "BIM Coordinator", "DDE", "Echipă BIM"],
            ["BIM-LIV-06", "Raport clash detection #1 – probleme identificate", "BCF + PDF", "BIM Coordinator", "DDE – Săpt. 2", "Proiectanți"],
            ["BIM-LIV-07", "Raport clash detection #2 – probleme rezolvate", "BCF + PDF", "BIM Coordinator", "DDE – Săpt. 4", "Client"],
            ["BIM-LIV-08", "Model 4D (progres planificat – Navisworks)", "NWD + XLSX", "Contractor BIM Lead", "Start execuție", "Client, Diriginte"],
            ["BIM-LIV-09", "Rapoarte lunare progres BIM (model 4D actualizat)", "PDF + NWD", "Contractor BIM Lead", "Lunar", "Client"],
            ["BIM-LIV-10", "Raport cantități Takeoff vs. proiect", "PDF + XLSX", "Contractor", "Lunar", "Client"],
            ["BIM-LIV-11", "Model As-Built complet (toate disciplinele)", "RVT + IFC", "Contractor BIM Lead", "Recepție prelim.", "Client"],
            ["BIM-LIV-12", "Planuri As-Built (DWG georeferențiat)", "DWG + PDF", "Contractor BIM Lead", "Recepție prelim.", "Client"],
            ["BIM-LIV-13", "Carte tehnică digitală (COBie / PDF structurat)", "XLSX + PDF", "BIM Manager", "Recepție finală", "Client (arhivă)"],
            ["BIM-LIV-14", "BEP Final – actualizat cu toate modificările", "DOCX/PDF", "BIM Manager", "Recepție finală", "Client"],
        ]
    )

    add_heading(doc, "8.2 Jaloane BIM (Milestones)", 2)
    add_matrix_table(doc,
        ["Jalonul", "Eveniment", "Livrabile asociate"],
        [
            ["M0", "Semnare contract / Kick-off proiect", "BIM-LIV-01 (BEP v1.0)"],
            ["M1", "Finalizare Proiect Tehnic (PT)", "BIM-LIV-02, 03, 04"],
            ["M2", "Finalizare Detalii de Execuție (DDE)", "BIM-LIV-05, 06, 07"],
            ["M3", "Predare șantier – Start execuție", "BIM-LIV-08"],
            ["M4", "Progres execuție 30%", "BIM-LIV-09, 10"],
            ["M5", "Progres execuție 60%", "BIM-LIV-09, 10"],
            ["M6", "Progres execuție 90%", "BIM-LIV-09, 10"],
            ["M7", "Finalizare lucrări – Recepție preliminară (HG 273/94)", "BIM-LIV-11, 12"],
            ["M8", "Recepție finală (după expirare perioadă garanție)", "BIM-LIV-13, 14"],
        ]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 9. COORDONARE BIM ȘI CLASH DETECTION
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. COORDONARE BIM ȘI CLASH DETECTION", 1, BLUE_DARK)

    add_heading(doc, "9.1 Procesul de coordonare", 2)
    add_body(doc,
        "Coordonarea modelelor multi-disciplinare se desfășoară prin cicluri săptămânale "
        "de verificare a interferențelor (clash detection) în Autodesk Navisworks Manage "
        "și/sau ACC Coordinate (cloud)."
    )
    add_body(doc, "Fluxul de coordonare:", bold=True)
    add_bullet(doc, "1. BIM Authorii publică modele disciplinare în zona Shared a CDE (format IFC/NWC)")
    add_bullet(doc, "2. BIM Coordinator federează modelele și rulează clash detection")
    add_bullet(doc, "3. Problemele identificate se exportă ca raport BCF 2.1 și se atribuie responsabililor")
    add_bullet(doc, "4. Proiectanții rezolvă interferențele și republică modelele corectate")
    add_bullet(doc, "5. BIM Manager verifică rezolvarea și aprobă trecerea în zona Published")
    add_bullet(doc, "6. Procesul se repetă până la zero interferențe critice înainte de DDE")

    add_heading(doc, "9.2 Tipuri de interferențe verificate", 2)
    add_matrix_table(doc,
        ["Tip Clash", "Descriere", "Prioritate", "Termen remediere"],
        [
            ["Hard Clash", "Intersecție fizică între elemente (ex. conductă vs. dig)", "CRITIC", "3 zile lucrătoare"],
            ["Soft Clash / Clearance", "Distanță insuficientă față de cerințe constructive", "MAJOR", "5 zile lucrătoare"],
            ["Workflow Clash", "Secvența de execuție incorectă (ex. etanșare înainte de profilare)", "MINOR", "10 zile lucrătoare"],
            ["Duplicat", "Elemente duplicate între modele diferite", "INFO", "La următoarea revizie"],
        ]
    )

    add_heading(doc, "9.3 Ședințe BIM (BIM Coordination Meetings)", 2)
    add_body(doc, "Se vor organiza ședințe BIM periodice cu următoarele frecvențe:")
    add_bullet(doc, "Proiectare (faza PT/DDE): săptămânal (online prin ACC / Teams), ~60 min")
    add_bullet(doc, "Execuție: bi-săptămânal (online), ~45 min")
    add_bullet(doc, "Milestone review (M1–M7): față în față / hibrid, cu participarea clientului")
    add_body(doc,
        "\nMinutele ședințelor BIM se înregistrează ca documente în CDE (folderul 08_Comunicări) "
        "în format PDF, semnate de BIM Manager."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 10. DOCUMENTAȚIE AS-BUILT ȘI PREDARE
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10. DOCUMENTAȚIE AS-BUILT ȘI PREDARE", 1, BLUE_DARK)

    add_heading(doc, "10.1 Cerințe As-Built", 2)
    add_body(doc,
        "Modelul As-Built reprezintă starea reală a construcției la momentul recepției "
        "preliminare (Art. 16 din CS, conform HG 273/94). Acesta trebuie să reflecte "
        "orice deviere față de proiect, autorizată în prealabil prin RFI și/sau dispoziție "
        "de șantier."
    )
    add_body(doc, "Cerințe minimale pentru modelul As-Built:", bold=True)
    add_bullet(doc, "LOD 400 pentru toate elementele principale (etanșare, drenaj, diguri, rampe)")
    add_bullet(doc, "Coordonate topografice reale, verificate prin ridicare topo post-execuție")
    add_bullet(doc, "Specificații complete materiale: producător, lot, certificate conformitate, PV recepție calitativă")
    add_bullet(doc, "Fotografii documentare asociate elementelor de model (linkate în ACC Docs)")
    add_bullet(doc, "Procese verbale de lucrări ascunse atașate elementelor relevante din model")

    add_heading(doc, "10.2 Dosarul Cărții Tehnice (digital)", 2)
    add_body(doc,
        "Cartea tehnică a construcției se va întocmi digital, în conformitate cu Legea 10/1995 "
        "și HG 343/2017, și va fi arhivată în CDE în folderul 07_Recepție. Va conține:"
    )
    add_bullet(doc, "A. Documentația de bază: proiectul as-built complet (planuri, modele IFC, memorii)")
    add_bullet(doc, "B. Documente de execuție: procese verbale lucrări ascunse, PV recepție faze, teste calitate")
    add_bullet(doc, "C. Documente de recepție: PV recepție preliminară, PV recepție finală")
    add_bullet(doc, "D. Documente exploatare: instrucțiuni exploatare, plan monitorizare levigat, fișe echipamente")
    add_bullet(doc, "E. Modificări în timp: dispoziții de șantier, acte adiționale, RFI-uri aprobate")

    add_heading(doc, "10.3 Predare la beneficiar (CJ Timiș)", 2)
    add_body(doc,
        "La recepția preliminară, Contractor BIM Lead (Inoveco) și BIM Manager (COPLAN CAI SRL) "
        "predau beneficiarului:"
    )
    add_bullet(doc, "Acces la proiectul ACC – dosar complet, inclusiv modele, documente, fotografii")
    add_bullet(doc, "Export IFC 4.3 al modelului As-Built (pe suport fizic – HDD + cloud)")
    add_bullet(doc, "Export DWG/PDF al planurilor As-Built georeferențiate (format A1/A0)")
    add_bullet(doc, "Export COBie (XLSX) cu toate datele activelor (echipamente, materiale, garanții)")
    add_bullet(doc, "Manual de utilizare a CDE (ACC) pentru beneficiar – gestionare arhivă")
    add_bullet(doc, "Sesiune de training pentru personalul beneficiarului (~2-4 ore)")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 11. MANAGEMENTUL CALITĂȚII BIM
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11. MANAGEMENTUL CALITĂȚII BIM", 1, BLUE_DARK)

    add_heading(doc, "11.1 Verificări model (Model Quality Checks)", 2)
    add_body(doc, "Înainte de publicarea oricărui model în zona Shared a CDE, BIM Author-ul va efectua:")
    add_bullet(doc, "Verificare vizuală – nu există elemente vizibil eronate, duplicate sau în afara domeniului")
    add_bullet(doc, "Verificare LOD – elementele respectă nivelul de detaliu cerut pentru faza curentă")
    add_bullet(doc, "Verificare metadate – toate elementele au completate câmpurile obligatorii (tip, material, dimensiuni)")
    add_bullet(doc, "Verificare IFC export – modelul se exportă corect în IFC fără erori critice")
    add_bullet(doc, "Clash detection preliminar (self-check) – zero interferențe hard în propria disciplină")

    add_heading(doc, "11.2 Indicatori KPI BIM", 2)
    add_matrix_table(doc,
        ["KPI", "Țintă", "Frecvență măsurare", "Responsabil"],
        [
            ["Nr. clash-uri hard nerezolvate", "0 la emiterea DDE", "Săptămânal", "BIM Coordinator"],
            ["% modele publicate la termen (față de plan)", "≥ 90%", "Lunar", "BIM Manager"],
            ["Nr. RFI-uri nerezolvate > 10 zile", "0", "Săptămânal", "Contractor BIM Lead"],
            ["Conformitate LOD față de cerințe BEP", "100% elemente principale", "Per jalonul M1-M7", "BIM Manager"],
            ["% documente indexate corect în CDE", "100%", "Lunar", "CDE Admin"],
            ["Timp mediu rezolvare clash critic", "≤ 3 zile lucrătoare", "Per sprint coordonare", "BIM Coordinator"],
        ]
    )

    add_heading(doc, "11.3 Audituri BIM", 2)
    add_body(doc, "Se vor efectua audituri BIM la următoarele momente:")
    add_bullet(doc, "Audit intern – la fiecare jalonul major (M1-M7); efectuat de BIM Manager")
    add_bullet(doc, "Audit extern – la recepția preliminară; efectuat de reprezentantul clientului sau o firmă terță")
    add_body(doc,
        "\nRaportul de audit va evalua conformitatea modelelor cu BEP-ul, acuratețea documentației "
        "în CDE și calitatea predării As-Built."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 12. PLAN DE IMPLEMENTARE BIM
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "12. PLAN DE IMPLEMENTARE BIM", 1, BLUE_DARK)

    add_heading(doc, "12.1 Faza 1 – Pregătire și configurare (M0)", 2)
    add_bullet(doc, "Semnare BEP de toți membrii echipei BIM")
    add_bullet(doc, "Creare și configurare proiect în ACC (structură foldere, permisiuni utilizatori)")
    add_bullet(doc, "Distribuire template-uri Revit, standarde de modelare, convenții denumire")
    add_bullet(doc, "Instruire echipă (kick-off BIM meeting) – prezentare BEP, ACC, fluxuri")
    add_bullet(doc, "Stabilire punct de origine comun și sistem de coordonate georeferențiat")

    add_heading(doc, "12.2 Faza 2 – Proiectare PT/DDE (M0 → M2)", 2)
    add_bullet(doc, "Modelare 3D Revit pe discipline (structuri, civil, instalații drenaj)")
    add_bullet(doc, "Cicluri săptămânale de coordonare și clash detection (Navisworks / ACC Coordinate)")
    add_bullet(doc, "Extragere cantități terasamente (Civil 3D volume calculation)")
    add_bullet(doc, "Publicare modele IFC în CDE la fiecare revizie majoră")
    add_bullet(doc, "Emitere BIM-LIV-01 până la BIM-LIV-07 conform calendarului")

    add_heading(doc, "12.3 Faza 3 – Execuție (M3 → M6)", 2)
    add_bullet(doc, "Conectare model Revit la programul de lucrări (4D BIM în Navisworks)")
    add_bullet(doc, "Utilizare ACC Build pentru RFI-uri, submittal-uri, inspecții și punch lists")
    add_bullet(doc, "Utilizare ACC Takeoff pentru verificare cantități realizate vs. proiect")
    add_bullet(doc, "Actualizare progresivă model As-Built (devieri față de proiect documentate)")
    add_bullet(doc, "Documentare foto lucrări ascunse (etanșare, geosintétice) în ACC Docs")
    add_bullet(doc, "Rapoarte lunare progres BIM (BIM-LIV-09, 10)")

    add_heading(doc, "12.4 Faza 4 – Finalizare și predare (M7 → M8)", 2)
    add_bullet(doc, "Finalizare model As-Built complet (LOD 400 toate disciplinele)")
    add_bullet(doc, "Ridicare topografică post-execuție și actualizare model cu coordonate reale")
    add_bullet(doc, "Compilare carte tehnică digitală (COBie export, PDF structurat)")
    add_bullet(doc, "Predare către CJ Timiș: modele IFC, DWG As-Built, PDF, XLSX COBie")
    add_bullet(doc, "Training beneficiar pentru utilizarea arhivei ACC")
    add_bullet(doc, "Emitere BEP Final (BIM-LIV-14) și arhivare proiect")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════════
    # 13. ANEXE ȘI APROBĂRI
    # ════════════════════════════════════════════════════════════════════════════
    add_heading(doc, "13. APROBĂRI BEP", 1, BLUE_DARK)

    add_body(doc,
        "Prezentul BEP este supus aprobării tuturor părților implicate. Semnătura electronică "
        "sau umedă atestă acceptarea cerințelor și angajamentul față de implementarea lor."
    )
    add_matrix_table(doc,
        ["Rol", "Nume / Organizație", "Funcție", "Data", "Semnătură"],
        [
            ["Client / Beneficiar",   "CJ Timiș",            "Reprezentant legal",    "_________", "___________"],
            ["BIM Manager",           "COPLAN CAI SRL",       "BIM Manager desemnat",  "_________", "___________"],
            ["BIM Coordinator Civil", "COPLAN CAI SRL",       "Proiectant civil",      "_________", "___________"],
            ["Contractor BIM Lead",   "Inoveco SRL",          "BIM Lead șantier",      "_________", "___________"],
            ["Diriginte de Șantier",  "Desemnat de CJ Timiș", "Diriginte autorizat",   "_________", "___________"],
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "Istoricul reviziilor BEP", 2)
    add_matrix_table(doc,
        ["Versiune", "Data", "Modificări", "Aprobat de"],
        [
            ["v1.0", TODAY, "Emitere inițială – Draft pentru revizuire și aprobare", "BIM Manager"],
            ["v1.1", "TBD", "Actualizare după feedback client și contractor", "BIM Manager + Client"],
            ["v2.0", "TBD", "Emitere la startul execuției – versiune completă", "BIM Manager + Client"],
        ]
    )

    doc.add_paragraph()
    add_body(doc,
        "Notă: Prezentul BEP va fi revizuit obligatoriu la fiecare jalonul major (M0–M8) "
        "sau ori de câte ori apar modificări semnificative în proiect. Versiunea curentă "
        "aprobată este întotdeauna disponibilă în CDE (ACC Docs), folderul 00_BEP și Standarde.",
        italic=True
    )

    # ── Salvare ─────────────────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f"\n✓ BEP generat cu succes: {OUTPUT_FILE}")
    print(f"  Deschide documentul cu Microsoft Word sau LibreOffice.")


if __name__ == "__main__":
    main()
